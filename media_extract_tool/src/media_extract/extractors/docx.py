from __future__ import annotations

import hashlib
from importlib.metadata import PackageNotFoundError, version
from pathlib import Path
from typing import Any

import docx

from media_extract.udr import new_block_id, now_iso


def _docx_version() -> str:
    try:
        return version("python-docx")
    except PackageNotFoundError:
        return getattr(docx, "__version__", "unknown")


def _heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    s = style_name.strip()
    for prefix in ("Heading ", "标题 "):
        if s.startswith(prefix):
            rest = s[len(prefix) :].strip()
            if rest.isdigit():
                lv = int(rest)
                if 1 <= lv <= 9:
                    return lv
    if s.lower().startswith("heading"):
        tail = s[7:].strip()
        if tail.isdigit():
            lv = int(tail)
            if 1 <= lv <= 9:
                return lv
    return None


def _is_list_paragraph(para: Any) -> bool:
    try:
        p = para._p
        if p.pPr is not None and p.pPr.numPr is not None:
            return True
    except Exception:
        pass
    style = para.style.name if para.style else ""
    return "list" in (style or "").lower()


def extract_docx(path: Path, source_uri: str) -> dict[str, Any]:
    raw = path.read_bytes()
    fp = hashlib.sha256(raw).hexdigest()
    document: dict[str, Any] = {
        "sourceFormat": "docx",
        "sourceUri": source_uri,
        "fingerprint": fp,
        "originalFilename": path.name,
        "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "byteSize": len(raw),
        "ingestedAt": now_iso(),
        "pipeline": {
            "libraryVersions": {"python-docx": _docx_version()},
        },
    }

    d = docx.Document(str(path))
    blocks: list[dict[str, Any]] = []
    n = 0

    def push_block(
        btype: str,
        text: str,
        *,
        level: int | None = None,
        docx_path: str | None = None,
    ) -> None:
        nonlocal n
        n += 1
        bid = new_block_id(n)
        o: dict[str, Any] = {
            "blockId": bid,
            "type": btype,
            "text": text,
            "confidence": None,
        }
        if level is not None:
            o["level"] = level
        if docx_path:
            o["location"] = {"docxPath": docx_path}
        blocks.append(o)

    for i, para in enumerate(d.paragraphs):
        text = (para.text or "").strip()
        style = para.style.name if para.style else None
        hl = _heading_level(style)
        loc = f"paragraph[{i}]"
        if hl is not None and text:
            push_block("heading", text, level=hl, docx_path=loc)
        elif text:
            if _is_list_paragraph(para):
                push_block("list_item", text, docx_path=loc)
            else:
                push_block("paragraph", text, docx_path=loc)

    for ti, table in enumerate(d.tables):
        push_block("table", "", docx_path=f"table[{ti}]")
        for ri, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            row_text = " | ".join(cells)
            push_block("table_row", row_text, docx_path=f"table[{ti}]/row[{ri}]")
            for ci, cell in enumerate(row.cells):
                ct = cell.text.strip()
                if ct:
                    push_block("table_cell", ct, docx_path=f"table[{ti}]/r{ri}c{ci}")

    return {"schemaVersion": "1.0.0", "document": document, "blocks": blocks, "assets": []}
